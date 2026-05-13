from flask import Flask, render_template, request, url_for, redirect, flash, session
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from flask_wtf import FlaskForm
from wtforms import StringField, PasswordField, SubmitField
from wtforms.validators import DataRequired, Email, EqualTo, Length
from werkzeug.security import generate_password_hash, check_password_hash
import os
import json
import hashlib
from datetime import datetime
from PIL import Image, ExifTags
import PyPDF2
import docx
from io import StringIO


# ==================== APP CONFIGURATION ====================
app = Flask(__name__)

app.config['SECRET_KEY'] = 'your-secret-key-2024'

# Database Configuration (Production vs Local)
if os.environ.get('DATABASE_URL'):
    database_url = os.environ.get('DATABASE_URL')
    if database_url and database_url.startswith('postgres://'):
        database_url = database_url.replace('postgres://', 'postgresql://', 1)
    app.config['SQLALCHEMY_DATABASE_URI'] = database_url
else:
    app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///database.db'

app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max

# Allowed file types
ALLOWED_EXTENSIONS = {
    'png', 'jpg', 'jpeg', 'gif', 'bmp',  # Images
    'pdf',                                 # PDF Documents
    'docx',                                # Word Documents
    'txt'                                  # Text files
}

db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'
login_manager.login_message_category = 'error'
login_manager.session_protection = "strong"

if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])


# ==================== DATABASE MODELS ====================

class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password_hash = db.Column(db.String(200), nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    def set_password(self, password):
        self.password_hash = generate_password_hash(password, method='pbkdf2:sha256')

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)


class Analysis(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    filename = db.Column(db.String(200), nullable=False)
    file_hash = db.Column(db.String(128), nullable=False)
    suspicion_score = db.Column(db.Integer, default=0)
    status = db.Column(db.String(50), default='CLEAN')
    forensic_findings = db.Column(db.Text, default='[]')
    file_path = db.Column(db.String(500))
    analysis_date = db.Column(db.DateTime, default=datetime.utcnow)

    def get_findings_list(self):
        try:
            return json.loads(self.forensic_findings)
        except:
            return []

    def set_findings_list(self, findings_list):
        self.forensic_findings = json.dumps(findings_list)


# ==================== FORMS ====================

class RegistrationForm(FlaskForm):
    username = StringField('Username', validators=[DataRequired(), Length(min=3, max=80)])
    email = StringField('Email', validators=[DataRequired(), Email()])
    password = PasswordField('Password', validators=[DataRequired(), Length(min=6)])
    confirm_password = PasswordField('Confirm Password', validators=[DataRequired(), EqualTo('password')])
    submit = SubmitField('Register')


class LoginForm(FlaskForm):
    username = StringField('Username', validators=[DataRequired()])
    password = PasswordField('Password', validators=[DataRequired()])
    submit = SubmitField('Login')


# ==================== AUTHENTICATION HELPERS ====================

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


def generate_csrf_token():
    if '_csrf_token' not in session:
        from secrets import token_hex
        session['_csrf_token'] = token_hex(16)
    return session['_csrf_token']


app.jinja_env.globals['csrf_token'] = generate_csrf_token


# ==================== FILE HASHING ====================

def generate_file_hash(file_path):
    """Generate SHA-512 hash of any file"""
    sha512_hash = hashlib.sha512()
    try:
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha512_hash.update(byte_block)
        return sha512_hash.hexdigest()
    except Exception as e:
        return f"Error: {str(e)}"


# ==================== PDF ANALYSIS ====================

def analyze_pdf(file_path):
    """Analyze PDF files for suspicious content"""
    score = 0
    findings = []

    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)

            findings.append(f"PDF has {num_pages} page(s)")

            if num_pages == 0:
                score += 30
                findings.append("⚠️ PDF has no pages - Possible corruption")
            elif num_pages == 1:
                score += 5
                findings.append("Single page PDF document")

            metadata = pdf_reader.metadata
            if metadata:
                findings.append("PDF metadata found")
                if metadata.get('/Author'):
                    findings.append(f"   Author: {metadata.get('/Author')}")
                if metadata.get('/Producer'):
                    producer = metadata.get('/Producer', '').lower()
                    if 'photoshop' in producer or 'editor' in producer:
                        score += 20
                        findings.append(f"⚠️ Created with editing software: {metadata.get('/Producer')}")
            else:
                score += 25
                findings.append("⚠️ No metadata found - File may be stripped")

            file_size = os.path.getsize(file_path)
            size_mb = file_size / (1024 * 1024)
            if size_mb > 10:
                score += 10
                findings.append(f"⚠️ Large PDF ({size_mb:.1f}MB) - Check for hidden content")

    except Exception as e:
        score += 20
        findings.append(f"⚠️ Could not analyze PDF: {str(e)}")

    return score, findings


# ==================== DOCX ANALYSIS ====================

def analyze_docx(file_path):
    """Analyze Word documents for suspicious content"""
    score = 0
    findings = []

    try:
        doc = docx.Document(file_path)
        num_paragraphs = len(doc.paragraphs)
        findings.append(f"Word document has {num_paragraphs} paragraphs")

        core_props = doc.core_properties
        if core_props.author:
            findings.append(f"   Author: {core_props.author}")
        if core_props.last_modified_by:
            findings.append(f"   Last modified by: {core_props.last_modified_by}")

        full_text = ' '.join([p.text for p in doc.paragraphs[:50]])
        suspicious_words = ['confidential', 'secret', 'internal', 'draft', 'edit', 'modified']
        found_suspicious = []
        for word in suspicious_words:
            if word.lower() in full_text.lower():
                found_suspicious.append(word)

        if found_suspicious:
            score += 15
            findings.append(f"⚠️ Suspicious keywords found: {', '.join(found_suspicious)}")

        file_size = os.path.getsize(file_path)
        size_mb = file_size / (1024 * 1024)
        if size_mb > 5:
            score += 10
            findings.append(f"⚠️ Large document ({size_mb:.1f}MB)")

    except Exception as e:
        score += 20
        findings.append(f"⚠️ Could not analyze DOCX: {str(e)}")

    return score, findings


# ==================== TXT ANALYSIS ====================

def analyze_txt(file_path):
    """Analyze text files for suspicious patterns"""
    score = 0
    findings = []

    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read(10000)
            lines = content.count('\n') + 1
            words = len(content.split())

            findings.append(f"Text file: ~{words} words, {lines} lines")

            if len(content) < 50:
                score += 10
                findings.append("⚠️ Very short text file - Possibly incomplete")

            suspicious_patterns = {
                'password': 'Password mentioned',
                'confidential': 'Confidential document',
                'secret': 'Secret content detected',
                'edited': 'May have been edited',
                'modified': 'Modified document'
            }

            detected = []
            for pattern, message in suspicious_patterns.items():
                if pattern.lower() in content.lower():
                    detected.append(message)

            if detected:
                score += 15
                findings.append(f"⚠️ Suspicious content: {', '.join(detected)}")

            file_size = os.path.getsize(file_path)
            size_mb = file_size / (1024 * 1024)
            if size_mb > 1:
                score += 5
                findings.append(f"📁 Large text file ({size_mb:.1f}MB)")

    except Exception as e:
        score += 15
        findings.append(f"⚠️ Could not analyze text file: {str(e)}")

    return score, findings


# ==================== IMAGE ANALYSIS ====================

def analyze_image(file_path):
    """Analyze image files for tampering"""
    score = 0
    findings = []

    try:
        with Image.open(file_path) as img:
            img_format = img.format
            findings.append(f"Image format: {img_format}")

            width, height = img.size
            if width < 100 or height < 100:
                score += 15
                findings.append(f"⚠️ Very small dimensions ({width}x{height})")
            else:
                findings.append(f"Dimensions: {width}x{height}")

            exif_data = img._getexif()

            if exif_data is None or len(exif_data) == 0:
                score += 30
                findings.append("⚠️ NO METADATA FOUND - File may be edited or from social media")
            else:
                findings.append(f"✅ Metadata found ({len(exif_data)} tags)")

                software_detected = False
                for tag_id, value in exif_data.items():
                    tag_name = ExifTags.TAGS.get(tag_id, str(tag_id))

                    if 'Software' in tag_name:
                        software = str(value).lower()
                        software_detected = True
                        if 'photoshop' in software:
                            score += 35
                            findings.append("⚠️ PHOTOSHOP DETECTED - High suspicion!")
                        elif 'canva' in software:
                            score += 25
                            findings.append("⚠️ CANVA DETECTED - Image likely edited")
                        elif 'gimp' in software:
                            score += 20
                            findings.append("⚠️ GIMP DETECTED")
                        else:
                            score += 15
                            findings.append(f"⚠️ Editing software: {value}")

                    if 'Make' in tag_name or 'Model' in tag_name:
                        findings.append(f"📱 Camera: {value}")
                        score = max(0, score - 5)

                if not software_detected:
                    findings.append("✅ No editing software detected")

            file_size = os.path.getsize(file_path)
            size_kb = file_size / 1024
            if size_kb < 10:
                score += 15
                findings.append(f"⚠️ Tiny file ({size_kb:.1f}KB)")

    except Exception as e:
        score += 20
        findings.append(f"⚠️ Image analysis error: {str(e)}")

    return score, findings


# ==================== MAIN FORENSIC FUNCTION ====================

def analyze_file_forensic(file_path):
    """Main forensic analysis based on file type"""
    filename = os.path.basename(file_path).lower()
    file_ext = filename.split('.')[-1] if '.' in filename else ''

    score = 0
    findings = []

    # Filename check
    suspicious_patterns = ['edit', 'copy', 'modified', 'screenshot', 'whatsapp', 'instagram', 'temp']
    for pattern in suspicious_patterns:
        if pattern in filename:
            score += 10
            findings.append(f"⚠️ Suspicious filename: '{pattern}' detected")
            break

    # File size check
    file_size = os.path.getsize(file_path)
    size_mb = file_size / (1024 * 1024)
    if size_mb == 0:
        score += 30
        findings.append("⚠️ File is empty!")
    elif size_mb > 20:
        score += 10
        findings.append(f"Large file ({size_mb:.1f}MB) - Check thoroughly")

    # Analyze based on file type
    if file_ext in ['png', 'jpg', 'jpeg', 'gif', 'bmp']:
        img_score, img_findings = analyze_image(file_path)
        score += img_score
        findings.extend(img_findings)

    elif file_ext == 'pdf':
        pdf_score, pdf_findings = analyze_pdf(file_path)
        score += pdf_score
        findings.extend(pdf_findings)

    elif file_ext == 'docx':
        docx_score, docx_findings = analyze_docx(file_path)
        score += docx_score
        findings.extend(docx_findings)

    elif file_ext == 'txt':
        txt_score, txt_findings = analyze_txt(file_path)
        score += txt_score
        findings.extend(txt_findings)

    else:
        findings.append(f"⚠️ Unknown file type: {file_ext}")
        score += 20

    # Cap score at 100
    score = min(max(score, 0), 100)

    # Determine status
    if score >= 50:
        status = "SUSPICIOUS"
    elif score >= 25:
        status = "POSSIBLY EDITED"
    else:
        status = "CLEAN"

    return {
        "score": score,
        "findings": findings,
        "status": status
    }


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def save_analysis(user_id, filename, file_hash, suspicion_score, status, findings, file_path):
    analysis = Analysis(
        user_id=user_id,
        filename=filename,
        file_hash=file_hash,
        suspicion_score=suspicion_score,
        status=status,
        file_path=file_path
    )
    analysis.set_findings_list(findings)
    db.session.add(analysis)
    db.session.commit()
    return analysis


# ==================== ROUTES ====================

@app.route('/')
def index():
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))
    return redirect(url_for('login'))


@app.route('/register', methods=['GET', 'POST'])
def register():
    # If already logged in, go to dashboard
    if current_user.is_authenticated:
        return redirect(url_for('dashboard'))

    form = RegistrationForm()
    if form.validate_on_submit():
        if User.query.filter_by(username=form.username.data).first():
            flash('Username already exists', 'error')
            return render_template('register.html', form=form)

        if User.query.filter_by(email=form.email.data).first():
            flash('Email already registered', 'error')
            return render_template('register.html', form=form)

        user = User(username=form.username.data, email=form.email.data)
        user.set_password(form.password.data)

        db.session.add(user)
        db.session.commit()

        flash(f'Registration successful! Welcome {user.username}!', 'success')
        return redirect(url_for('login'))

    return render_template('register.html', form=form)


@app.route('/login', methods=['GET', 'POST'])
def login():
    # YEH BLOCK BILKUL HATA DO ↓
    # if current_user.is_authenticated:
    #     return redirect(url_for('dashboard'))
    
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(username=form.username.data).first()
        if user and user.check_password(form.password.data):
            login_user(user, remember=False)  # remember=True se False karo
            flash(f'Welcome back, {user.username}!', 'success')
            return redirect(url_for('dashboard'))
        else:
            flash('Invalid username or password', 'error')

    return render_template('login.html', form=form)


@app.route('/logout')
@login_required
def logout():
    logout_user()
    session.pop('_user_id', None)
    session.pop('_fresh', None)
    session.pop('_id', None)
    session.modified = True
    session.clear()
    flash('You have been logged out.', 'info')
    return redirect(url_for('login'))

@app.route('/dashboard')
@login_required
def dashboard():
    history = Analysis.query.filter_by(user_id=current_user.id).order_by(Analysis.analysis_date.desc()).all()
    return render_template('dashboard.html', history=history, current_user=current_user)


@app.route('/upload')
@login_required
def home():
    return render_template('upload.html', current_user=current_user)


@app.route('/upload', methods=['POST'])
@login_required
def upload_file():
    if 'file' not in request.files:
        flash('No file selected', 'error')
        return redirect(url_for('home'))

    if request.form.get('_csrf_token') != session.get('_csrf_token'):
        flash('CSRF token validation failed', 'error')
        return redirect(url_for('home'))

    file = request.files['file']

    if file.filename == '':
        flash('No file selected', 'error')
        return redirect(url_for('home'))

    safe_filename = file.filename.replace('/', '_').replace('\\', '_')

    if file and allowed_file(file.filename):
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_')
        unique_filename = timestamp + safe_filename
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
        file.save(file_path)

        file_hash = generate_file_hash(file_path)
        analysis_results = analyze_file_forensic(file_path)

        save_analysis(
            user_id=current_user.id,
            filename=safe_filename,
            file_hash=file_hash,
            suspicion_score=analysis_results.get('score', 0),
            status=analysis_results.get('status', 'CLEAN'),
            findings=analysis_results.get('findings', []),
            file_path=file_path
        )

        return render_template('report.html',
                               filename=safe_filename,
                               file_hash=file_hash,
                               forensic_findings=analysis_results.get('findings', []),
                               suspicion_score=analysis_results.get('score', 0),
                               status=analysis_results.get('status', 'Unknown'))
    else:
        flash('Security Error: File type not allowed! Allowed: PNG, JPG, JPEG, PDF, DOCX, TXT', 'error')
        return redirect(url_for('home'))


@app.route('/view_analysis/<int:analysis_id>')
@login_required
def view_analysis(analysis_id):
    analysis = Analysis.query.get_or_404(analysis_id)
    if analysis.user_id != current_user.id:
        flash('Access denied', 'error')
        return redirect(url_for('dashboard'))

    return render_template('report.html',
                           filename=analysis.filename,
                           file_hash=analysis.file_hash,
                           forensic_findings=analysis.get_findings_list(),
                           suspicion_score=analysis.suspicion_score,
                           status=analysis.status)


if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(debug=True)
